"""
Document Engine — extract text from PDF, DOCX, CSV files.

Preserves metadata (filename, page numbers, headings, tables where possible).
Produces clean text ready for instruction dataset generation.
"""

import json, re, csv, io, hashlib
from pathlib import Path
from datetime import datetime, timezone
from typing import List, Dict, Any, Optional
from app.core import settings, get_logger

logger = get_logger("app.engines.dataset.document_engine")


class DocumentChunk:
    """A chunk of extracted text with source metadata."""
    def __init__(self, text: str, source: str, page: int = 0, section: str = "",
                 chunk_type: str = "paragraph"):
        self.text = text.strip()
        self.source = source
        self.page = page
        self.section = section
        self.chunk_type = chunk_type
        self.id = hashlib.md5(f"{source}:{page}:{text[:80]}".encode()).hexdigest()[:12]

    def to_dict(self) -> Dict[str, Any]:
        return {"id": self.id, "text": self.text, "source": self.source,
                "page": self.page, "section": self.section or "body",
                "type": self.chunk_type}


class DocumentExtractor:
    """Extracts text from supported document formats."""

    SUPPORTED = {".pdf", ".docx", ".csv", ".txt", ".md"}

    def extract(self, file_path: Path) -> List[DocumentChunk]:
        """Extract text chunks from a document file."""
        suffix = file_path.suffix.lower()
        if suffix not in self.SUPPORTED:
            raise ValueError(f"Unsupported format: {suffix}. Supported: {self.SUPPORTED}")

        source = file_path.name
        if suffix == ".pdf":
            return self._extract_pdf(file_path, source)
        elif suffix == ".docx":
            return self._extract_docx(file_path, source)
        elif suffix == ".csv":
            return self._extract_csv(file_path, source)
        else:
            return self._extract_text(file_path, source)

    def extract_directory(self, directory: Path) -> List[DocumentChunk]:
        """Extract from all supported files in a directory."""
        chunks = []
        for f in sorted(directory.iterdir()):
            if f.suffix.lower() in self.SUPPORTED:
                try:
                    chunks.extend(self.extract(f))
                    logger.info("document_extracted", file=f.name, chunks=len(chunks))
                except Exception as e:
                    logger.error("extraction_failed", file=f.name, error=str(e))
        return chunks

    # ------------------------------------------------------------------
    # Format-specific extractors
    # ------------------------------------------------------------------

    def _extract_pdf(self, path: Path, source: str) -> List[DocumentChunk]:
        chunks = []
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(path))
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
                for para in paragraphs:
                    if len(para) > 20:
                        chunks.append(DocumentChunk(para, source, page=i + 1, chunk_type="paragraph"))
        except ImportError:
            logger.warning("PyPDF2 not installed. pip install PyPDF2")
        except Exception as e:
            logger.error("pdf_extraction_failed", file=source, error=str(e))
        return chunks

    def _extract_docx(self, path: Path, source: str) -> List[DocumentChunk]:
        chunks = []
        try:
            from docx import Document
            doc = Document(str(path))
            for para in doc.paragraphs:
                text = para.text.strip()
                if text and len(text) > 20:
                    heading = para.style.name if "Heading" in (para.style.name or "") else "body"
                    chunks.append(DocumentChunk(text, source, section=heading, chunk_type="paragraph"))
            # Extract tables
            for t_idx, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    table_text = "\n".join(rows)
                    chunks.append(DocumentChunk(table_text, source, section=f"table_{t_idx+1}", chunk_type="table"))
        except ImportError:
            logger.warning("python-docx not installed. pip install python-docx")
        except Exception as e:
            logger.error("docx_extraction_failed", file=source, error=str(e))
        return chunks

    def _extract_csv(self, path: Path, source: str) -> List[DocumentChunk]:
        chunks = []
        try:
            with open(path, "r", encoding="utf-8") as f:
                reader = csv.DictReader(f)
                rows = list(reader)
            if not rows:
                return chunks
            # Describe the CSV structure
            columns = list(rows[0].keys())
            description = f"This CSV file contains {len(rows)} rows with columns: {', '.join(columns)}. "
            sample_rows = rows[:5]
            description += "Sample data:\n"
            for i, row in enumerate(sample_rows):
                description += f"Row {i+1}: " + " | ".join(f"{k}: {v}" for k, v in row.items()) + "\n"
            chunks.append(DocumentChunk(description, source, section="csv_summary", chunk_type="structured"))
            # Individual row descriptions for detailed QA
            for i, row in enumerate(rows):
                row_text = json.dumps(row, ensure_ascii=False)
                chunks.append(DocumentChunk(row_text, source, page=i + 1, section="csv_row", chunk_type="row"))
        except Exception as e:
            logger.error("csv_extraction_failed", file=source, error=str(e))
        return chunks

    def _extract_text(self, path: Path, source: str) -> List[DocumentChunk]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip() and len(p.strip()) > 20]
        return [DocumentChunk(p, source, chunk_type="paragraph") for p in paragraphs]


# ------------------------------------------------------------------
# Instruction Dataset Generator
# ------------------------------------------------------------------

INSTRUCTION_TEMPLATES = [
    ("factual_qa", "According to the document, {context_hint}"),
    ("summary", "Summarize the information about {context_hint}."),
    ("definition", "What is {context_hint} as described in the document?"),
    ("explanation", "Explain {context_hint} based on the provided document."),
    ("policy", "What is the policy regarding {context_hint}?"),
    ("procedure", "What are the steps for {context_hint}?"),
    ("list", "List the key points about {context_hint} from the document."),
    ("comparison", "Compare {context_hint} as described in the document."),
    ("conversation", "Can you tell me about {context_hint} from the document?"),
]

class InstructionGenerator:
    """Generates diverse instruction-response pairs from document chunks.

    Creates natural QA pairs for supervised fine-tuning.
    Reuses Forge's existing JSONL schema: {"instruction": ..., "input": ..., "output": ...}
    """

    def __init__(self, max_pairs_per_chunk: int = 3):
        self._max_pairs = max_pairs_per_chunk
        self._extractor = DocumentExtractor()

    def generate_from_chunks(self, chunks: List[DocumentChunk]) -> List[Dict[str, str]]:
        """Generate instruction-response pairs from extracted chunks."""
        pairs = []
        for chunk in chunks:
            if len(chunk.text) < 30:
                continue
            # Generate 2-3 diverse questions per chunk
            num_questions = min(self._max_pairs, max(1, len(chunk.text) // 200))
            selected = [INSTRUCTION_TEMPLATES[i % len(INSTRUCTION_TEMPLATES)] for i in range(num_questions)]

            for qtype, template in selected:
                hint = self._extract_hint(chunk)
                instruction = template.format(context_hint=hint)
                pairs.append({
                    "instruction": instruction,
                    "input": f"Source: {chunk.source}" + (f" (page {chunk.page})" if chunk.page else ""),
                    "output": chunk.text[:2000],
                    "metadata": {"type": qtype, "source": chunk.source, "chunk_id": chunk.id},
                })
        return pairs

    def generate_from_files(self, paths: List[Path]) -> tuple:
        """Full pipeline: extract documents and generate instruction dataset.

        Returns (records, stats_dict).
        """
        all_chunks = []
        for p in paths:
            if p.is_dir():
                all_chunks.extend(self._extractor.extract_directory(p))
            elif p.is_file():
                all_chunks.extend(self._extractor.extract(p))

        if not all_chunks:
            return [], {"chunks": 0, "pairs": 0, "sources": []}

        pairs = self.generate_from_chunks(all_chunks)
        sources = list(set(c.source for c in all_chunks))

        return pairs, {
            "documents_processed": len(sources),
            "chunks_extracted": len(all_chunks),
            "pairs_generated": len(pairs),
            "sources": sources,
            "generated_at": datetime.now(timezone.utc).isoformat(),
        }

    def save_dataset(self, pairs: List[Dict], output_path: Path) -> Path:
        """Save instruction pairs as JSONL dataset for Forge training."""
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            for p in pairs:
                f.write(json.dumps({"instruction": p["instruction"], "input": p["input"],
                                     "output": p["output"]}, ensure_ascii=False) + "\n")
        return output_path

    @staticmethod
    def _extract_hint(chunk: DocumentChunk) -> str:
        """Extract a short topic hint from chunk text for question generation."""
        text = chunk.text
        # Use the first sentence or first 80 chars as the topic hint
        first_sentence = re.split(r'[.!?]', text)[0].strip()
        if 5 <= len(first_sentence) <= 120:
            return first_sentence
        return text[:80].rsplit(" ", 1)[0] + "..." if len(text) > 80 else text


extractor = DocumentExtractor()
generator = InstructionGenerator()
